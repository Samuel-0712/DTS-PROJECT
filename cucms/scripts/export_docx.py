import sqlite3
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

db_path = r"c:\Users\SAMUEL\Downloads\misc\cucms-frontend\cucms\cucms.db"
output_path = r"c:\Users\SAMUEL\Downloads\misc\cucms-frontend\cucms\database_visual_representation.docx"

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set padding/margins for table cells in twentieths of a point (dxa)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color):
    """Set cell background color (hex string)"""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def export_to_docx():
    if not os.path.exists(db_path):
        print(f"Error: Database file not found at {db_path}")
        return

    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # Get list of all tables
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%';")
    tables = [row[0] for row in cursor.fetchall()]

    # Sort tables in logical hierarchy
    order_pref = ['cafeteria', 'user', 'student', 'staff_member', 'prepaid_account', 'menu_item', 'order', 'order_item', 'payment', 'prepaid_transaction']
    tables = sorted(tables, key=lambda x: order_pref.index(x) if x in order_pref else len(order_pref))

    # Initialize Document
    doc = Document()

    # Page Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(24)
    title_run = title_p.add_run("TABLES (VISUAL REPRESENTATION OF DATABASE)")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.underline = True
    title_run.font.color.rgb = RGBColor(13, 27, 42) # Deep Navy

    for table_name in tables:
        # Heading for each table
        heading_p = doc.add_paragraph()
        heading_p.paragraph_format.space_before = Pt(18)
        heading_p.paragraph_format.space_after = Pt(8)
        heading_p.paragraph_format.keep_with_next = True
        
        heading_run = heading_p.add_run(table_name.capitalize())
        heading_run.font.name = 'Georgia'
        heading_run.font.size = Pt(16)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(13, 27, 42)

        # Get table schema (columns)
        cursor.execute(f'PRAGMA table_info("{table_name}");')
        columns_info = cursor.fetchall()
        headers = [col[1] for col in columns_info]

        # Get rows
        cursor.execute(f'SELECT * FROM "{table_name}";')
        rows = cursor.fetchall()

        if not rows:
            p = doc.add_paragraph()
            p.add_run("Table is empty.").font.italic = True
            doc.add_paragraph() # spacing
            continue

        # Create Word Table
        word_table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        word_table.style = 'Table Grid'
        word_table.autofit = True

        # Header formatting
        hdr_cells = word_table.rows[0].cells
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
            set_cell_shading(hdr_cells[i], "0D1B2A") # Deep Navy Header
            
            # Format text inside header
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255) # White text

        # Row formatting
        for row_idx, row_data in enumerate(rows):
            row_cells = word_table.rows[row_idx + 1].cells
            # Zebra striping color (light cream/gray)
            bg_color = "F7F4EE" if row_idx % 2 == 1 else "FFFFFF"
            
            for col_idx, val in enumerate(row_data):
                val_str = "NULL" if val is None else str(val)
                row_cells[col_idx].text = val_str
                set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=150, right=150)
                set_cell_shading(row_cells[col_idx], bg_color)
                
                # Format cell text
                p = row_cells[col_idx].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
                    if val is None:
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(122, 138, 156) # Muted gray for NULL
                    else:
                        run.font.color.rgb = RGBColor(13, 27, 42)

        # Spacing after table
        doc.add_paragraph().paragraph_format.space_before = Pt(12)

    conn.close()
    
    # Save the Word Document
    doc.save(output_path)
    print(f"Successfully exported visual representation Word Document to {output_path}")

if __name__ == "__main__":
    export_to_docx()
